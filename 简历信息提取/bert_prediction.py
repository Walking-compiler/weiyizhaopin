import json
from docx import Document
import pdfplumber
from transformers import BertTokenizer, BertForSequenceClassification, AdamW
from tqdm import tqdm
import os
#os.environ["CUDA_VISIBLE_DEVICES"] = "6"
def find_max_index(lst):
    max_value = max(lst)
    max_index = lst.index(max_value)
    return max_index

def top_n_max_indices(lst, n):
    # 获取列表中所有元素的索引和值
    index_value_pairs = list(enumerate(lst))
    # 对索引和值进行排序，以值降序排列
    sorted_pairs = sorted(index_value_pairs, key=lambda x: x[1], reverse=True)
    # 提取前n个最大值的索引
    top_n_indices = [pair[0] for pair in sorted_pairs[:n]]
    return top_n_indices


def pdf_get_text(path):
    with pdfplumber.open(path) as pdf:
        page = pdf.pages[0]   # 第一页的信息
        text = page.extract_text()
    return text

def get_paragraphs_text(path):
    """
    获取所有段落的文本
    :param path: word路径
    :return: list类型，如：
        ['Test', 'hello world', ...]
    """
    document = Document(path)
    # 有的简历是表格式样的，因此，不仅需要提取正文，还要提取表格
    col_keys = [] # 获取列名
    col_values = [] # 获取列值
    index_num = 0
    # 表格提取中，需要添加一个去重机制
    fore_str = ""
    cell_text = ""
    for table in document.tables:
        for row_index,row in enumerate(table.rows):
            for col_index,cell in enumerate(row.cells):
                if fore_str != cell.text:
                    if index_num % 2==0:
                        col_keys.append(cell.text)
                    else:
                        col_values.append(cell.text)
                    fore_str = cell.text
                    index_num +=1
                    cell_text += cell.text + '\n'
    # 提取正文文本
    paragraphs_text = ""
    for paragraph in document.paragraphs:
        # 拼接一个list,包括段落的结构和内容
        paragraphs_text += paragraph.text + "\n"
    return cell_text, paragraphs_text


def bert_predict(person_path):
    model_path = './fine_tuned_model'  # 微调后模型路径
    tokenizer_path = './fine_tuned_model'  # 分词器路径

    model = BertForSequenceClassification.from_pretrained(model_path)
    tokenizer = BertTokenizer.from_pretrained(tokenizer_path)

    if 'pdf' in person_path:
        text=pdf_get_text(person_path)
    else:
        cell_text, paragraphs_text = get_paragraphs_text(person_path)
        text = cell_text + paragraphs_text


    if '博士' in text:
        education='博士'
    elif '研究生' in text or '硕士' in text:
        education='硕士'
    elif '本科' in text:
        education='本科'
    elif '专科' in text:
        education='专科'
    else:
        education='暂无'


    max_length=model.config.max_position_embeddings
    inputs = tokenizer(text, return_tensors='pt',max_length=max_length, truncation=True)  # 对文本进行分词和编码
    #print(max_length)
    input_ids = inputs['input_ids']
    #print(len(input_ids[0]))
    attention_mask = inputs['attention_mask']
    token_type_ids = inputs['token_type_ids']

    outputs = model(input_ids, attention_mask=attention_mask, token_type_ids=token_type_ids)
    logits = outputs[0]


    job_recommend=[]
    with open("job4.json",'r') as f:
        job=json.load(f)
        #a=find_max_index(list(logits[0]))
        a=top_n_max_indices(list(logits[0]),5)
        #print(a)
        for i in a:
            for k in job.keys():
                if job[k]==i:
                    job_recommend.append(k)
                    #print(k)
        #print(logits)
    #print(job_recommend)
    return job_recommend,education

import random
def Job_recommend(job1,job2):
    #job_recommend,education=bert_predict(person_path)
    #job_recommend1=kg(person_path)
    #jobaaa={'shixi':0.5,"java":0.7}

    #job1=['实习生', '销售工程师', 'IMC通信项目助理', '软件服务工程师', '档案数字化专员']
    #先去重，再一边选2一边选3
    #job2=['web前端工程师', '售后工程师', 'ipc工程师', '软件服务工程师', '档案数字化专员']
    job3=[]
    job4=[]
    for i in job1:
        if i in job2:
            job3.append(i)
        else:
            job4.append(i)

    for j in job2:
        if j not in job1:
            job4.append(j)


    num_elements_to_select = 5-len(job3)
    random_elements = random.sample(job4, num_elements_to_select)
    for i in random_elements:
        job3.append(i)
    print(job3)
    return job3




if __name__ == '__main__':
    job,edu=bert_predict('./person1/1 (2).pdf')
    print(job)
    #Job_recommend('./person1/1.docx')