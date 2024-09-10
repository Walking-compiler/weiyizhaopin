#针对用户第一次提交简历和对应职位进行匹配
import profile

from torch._C._profiler import ProfilerActivity
from tqdm import tqdm
import torch
import torch.nn as nn
import torch.optim as optim
from torch.utils import data
from docx import Document
import pdfplumber
import os
import warnings
import argparse
import numpy as np
# from sklearn import metrics
from model import Bert_BiLSTM_CRF
from transformers import AdamW, get_linear_schedule_with_warmup
from util import NerDataset, PadBatch, VOCAB, tokenizer, tag2idx, idx2tag
from py2neo import Graph, Node, Relationship, NodeMatcher
import jieba
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import warnings
import time

import re
def repp(resume_text):
    # 匹配时间段的正则表达式，支持起始和结束年月或年月日，考虑中文和常见分隔符
    pattern_phone = re.compile(r'\b1[3-9]\d{9}\b')

    pattern_location=re.compile(r'([\u4e00-\u9fa5]+省[\u4e00-\u9fa5]+市)')
    pattern_workexperience=re.compile(r'([\u4e00-\u9fa5]+有限公司)')
    pattern_project= re.compile(r'项目经验(.*?)证书情况', re.DOTALL)
    pattern_name = re.compile(r'姓名(.*?)性别', re.DOTALL)
    pattern_gender = re.compile(r'性别(.*?)出生日期', re.DOTALL)
    pattern_birthday = re.compile(r'出生日期(.*?)民族', re.DOTALL)
    pattern_job=re.compile(r'(职位：[\u4e00-\u9fa5]+)')


    # 匹配电子邮件地址的正则表达式
    pattern_email = re.compile(r'\b[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+\b')

    pattern_time = re.compile(r'(\d{4}年\d{1,2}月)-(\d{4}年\d{1,2}月)')
    pattern_time2=re.compile(r'(\d{4}/\d{1,2})-(\d{4}/\d{1,2})')

    # 匹配学历的正则表达式
    pattern_education = re.compile(r'(本科|硕士|研究生|博士|专科|学士)')

    # 匹配就读大学的正则表达式
    pattern_university = re.compile(r'([\u4e00-\u9fa5]+(大学|学院))')


    # 查找匹配项
    matches_time2=pattern_time2.findall(resume_text)
    matches_name=pattern_name.findall(resume_text)
    matches_gender=pattern_gender.findall(resume_text)
    matches_birthday=pattern_birthday.findall(resume_text)
    matches_job=pattern_job.findall(resume_text)
    matches_project=pattern_project.findall(resume_text)
    mathes_workexperience=pattern_workexperience.findall(resume_text)
    matches_location=pattern_location.findall(resume_text)
    matches_phone=pattern_phone.findall(resume_text)
    matches_email=pattern_email.findall(resume_text)
    matches_time = pattern_time.findall(resume_text)
    matches_education = pattern_education.findall(resume_text)
    matches_university = pattern_university.findall(resume_text)

    matches_time1=[]
    matches_time3=[]
    # 输出结果，并格式化时间段信息
    for match in matches_time:
        start_part = match[0] if match[0] else match[2]
        end_part = match[1] if match[1] else match[3]
        matches_time1.append(''.join(start_part+'-'+end_part))
    for match in matches_time2:
        start_part = match[0] if match[0] else match[2]
        end_part = match[1] if match[1] else match[3]
        matches_time3.append(''.join(start_part + '-' + end_part))
    matches_time2=matches_time3
        #print(f"找到的时间段: {start_part} 至 {end_part}")
    print(matches_education)
    print(matches_university)
    print(matches_time1)
    matches = []
    for i in range(len(matches_time1)):
        matches.append(''.join(matches_time1[i] + '\t' + matches_university[i][0] + '\t' + matches_education[i]))
    print(matches_time2)
    print(mathes_workexperience)
    print(matches_job)
    print(matches_gender)
    print(matches_name)
    print(matches_birthday)
    shixi_job=[]
    for i in range(len(mathes_workexperience)):
        shixi_job.append(''.join(matches_time2[i])+'\t'+mathes_workexperience[i]+'\t'+matches_job[i])
    print(shixi_job)
    return matches_phone[0],matches_email[0],matches,matches_location[0],shixi_job,matches_project,matches_name[0].split('\n')[0],matches_gender[0].split('\n')[0],matches_birthday[0].split('\n')[0]

def tokenize_zh(text):
    """
    使用jieba进行中文分词
    """
    return jieba.lcut(text)

def compute_similarity(text1, text2):
    """
    计算两段文本的相似度
    """
    # 创建一个TF-IDF Vectorizer对象，使用jieba进行分词
    vectorizer = TfidfVectorizer(tokenizer=tokenize_zh, stop_words=[])

    # 将文本转换为TF-IDF向量
    tfidf_matrix = vectorizer.fit_transform([text1, text2])

    # 计算余弦相似度
    similarity = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])

    return similarity[0][0]

def test(model, iterator, device):
    # model.eval()
    Y_hat = []
    #with torch.no_grad(),torch.autograd.profiler.profile(enabled=True, use_cuda=False, record_shapes=True, profile_memory=False) as prof:
    with torch.no_grad():
        t_1 = time.time()
        for i, batch in enumerate(iterator):
            x, y, z = batch
            x = x.to(device)
            z = z.to(device)
            y_hat = model(x, y, z, is_test=True)
            # Save prediction
            for j in y_hat:
                Y_hat.extend(j)
            # Save labels
        #print(f"model cost {time.time()-t_1}")
    y_pred = [idx2tag[i] for i in Y_hat]

    return y_pred

def pdf_get_text(path):
    with pdfplumber.open(path) as pdf:
        page = pdf.pages[0]   # 第一页的信息
        text = page.extract_text()
        page=pdf.pages[1]
        text=text+page.extract_text()
    return text

def get_paragraphs_text(path):
    """
    获取所有段落的文本
    :param path: word路径
    :return: list类型，如：
        ['Test', 'hello world', ...]
    """
    document = Document(path)
    # 有的简历是表格式样的，因此，不仅需要提取正文，还要提取表格
    col_keys = [] # 获取列名
    col_values = [] # 获取列值
    index_num = 0
    # 表格提取中，需要添加一个去重机制
    fore_str = ""
    cell_text = ""
    for table in document.tables:
        for row_index,row in enumerate(table.rows):
            for col_index,cell in enumerate(row.cells):
                if fore_str != cell.text:
                    if index_num % 2==0:
                        col_keys.append(cell.text)
                    else:
                        col_values.append(cell.text)
                    fore_str = cell.text
                    index_num +=1
                    cell_text += cell.text + '\n'
    # 提取正文文本
    paragraphs_text = ""
    for paragraph in document.paragraphs:
        # 拼接一个list,包括段落的结构和内容
        paragraphs_text += paragraph.text + "\n"
    return cell_text, paragraphs_text

def bratann2BIO_format(text,bio_path1):
    """
    text:巴中兴文置信号
    使用brat 等标注工具，标注完导出为 brat-ann 格式如下：
    T1 (这里是\t)行政区划 0 2  (这里是\t)巴中
    T2 行政区划 2 4  兴文
    T3 核心词 4 6 置信
    """
    label = ['O' for _ in range(len(text))]
    with open(bio_path1, 'w', encoding='utf-8') as f:
        length=len(list(text))
        i=0
        for t, l in zip(list(text), label):
            if t == '\t' or t==' ' or t=='\n' or t==''or i<int(length*0.3):
                i+=1
                continue
            line = ' '.join([t + '\t' + l])  # 可以使用空格代替
            # print(line)
            f.write(line)
            f.write('\n')
        # fstream.write('\n')

def init_model():
    _best_val_loss = 1e18
    _best_val_acc = 1e-18
    warnings.filterwarnings('ignore',
                            message="The parameter 'token_pattern' will not be used since 'tokenizer' is not None")
    device = 'cpu'
    #device = 'cuda' if torch.cuda.is_available() else 'cpu'
    #print(device)
    model = Bert_BiLSTM_CRF(tag2idx).to(device)
    model.eval()
    model.load_state_dict(torch.load('./bert_bilstm_crf/best9.pth'),False)
    return model,device

def person_information(person_path):
    person_document=os.path.dirname(person_path)
    if 'pdf' in person_path:
        text=pdf_get_text(person_path)
    else:
        cell_text, paragraphs_text = get_paragraphs_text(person_path)
        text = cell_text + paragraphs_text


    matchphone,matchemail,matches,matcheslocation,shixi_job,matchesproject,matchesname,matchesgender,matchesbithday=repp(text)
    #print(matchphone)
    #print(matchemail)
    #print(matches)

    with open(person_document+'/personinformation','w',encoding='utf-8') as f2:
        f2.write(repr(matchphone))
        f2.write('\n')
        f2.write(repr(matchemail))
        f2.write('\n')
        f2.write(repr(matches))

    return matchphone,matchemail,matches,matcheslocation[0],shixi_job,matchesproject,matchesname,matchesgender,matchesbithday



def personpredict(person_path,model,device):
    #t_1 = time.time()
    num = 1
    #person_path = './person1/1.docx'
    person_document=os.path.dirname(person_path)
    #print(person_document)
    person_predict_path = person_document+'/' + str(num) + '.txt'
    person_predict_bioprocess_path = person_document+'/' + str(num + 1) + '.txt'
    person_information_path=person_document+'/company.txt'
    bio_path1=person_document+'/train.txt'
    bio_path2=person_document+'/train1.txt'
    #job_desire = '前端开发工程师'


    if 'pdf' in person_path:
        text=pdf_get_text(person_path)
    else:
        cell_text, paragraphs_text = get_paragraphs_text(person_path)
        text = cell_text + paragraphs_text

    matchphone,matchemail,matches,matcheslocation,shixi_job,matchesproject,matchesname,matchesgender,matchesbithday= repp(text)
    # print(matchphone)
    # print(matchemail)
    # print(matches)
    '''
    with open(person_document + '/personinformation', 'w', encoding='utf-8') as f2:
        f2.write(repr(matchphone))
        f2.write('\n')
        f2.write(repr(matchemail))
        f2.write('\n')
        f2.write(repr(matches))
    '''
    if '博士' in text:
        education='博士'
    elif '研究生' in text or '硕士' in text:
        education='硕士'
    elif '本科' in text:
        education='本科'
    elif '专科' in text:
        education='专科'
    else:
        education='暂无'

    bratann2BIO_format(text,bio_path1)
    parser = argparse.ArgumentParser()
    parser.add_argument("--testset", type=str, default=bio_path1)
    parser.add_argument("--batch_size", type=int, default=64)
    ner = parser.parse_args()


    test_dataset = NerDataset(ner.testset,bio_path2=bio_path2)

    test_iter = data.DataLoader(dataset=test_dataset,
                                batch_size=(ner.batch_size) // 2,
                                shuffle=False,
                                num_workers=0,
                                collate_fn=PadBatch
                                )
    #print(f"before forward cost {time.time()-t_1}")
    #对简历进行BIO命名实体识别
    y_pred = test(model, test_iter, device)

    #print(f"afer model cost {time.time()-t_1}")
    #以规定格式保存为BIO文件
    with open(bio_path2, 'r+', encoding='utf-8') as f:
        with open(person_predict_path, 'w', encoding='utf-8') as f1:
            j = 0
            for i in f:
                while True:
                    if y_pred[j] != '[CLS]' and y_pred[j] != '[SEP]':
                        line = ' '.join([i.split('\n')[0] + '\t' + y_pred[j]])
                        f1.write(line)
                        f1.write('\n')
                        j += 1
                        break
                    else:
                        j += 1



    #bio整合
    with open(person_predict_path, 'r+', encoding='utf-8') as f:
        JN = []
        JS = []
        JR = []
        PS = []
        words = []
        i = 0
        t = list(f)
        # print(t)
        # print(len(t))
        while i < len(t):
            word = t[i].split('\t')[0]
            tag = t[i].split('\t')[1]
            # print(tag)
            if tag == 'B-JN\n':
                words.append(word)
                i += 1
                if i >= len(t): break
                word = t[i].split('\t')[0]
                tag = t[i].split('\t')[1]
                while tag == 'I-JN\n':
                    words.append(word)
                    i += 1
                    if i >= len(t): break
                    word = t[i].split('\t')[0]
                    tag = t[i].split('\t')[1]
                if ''.join(map(str, words)).lower() not in JN:
                    JN.append(''.join(map(str, words)).lower())
                    words = []
            elif tag == 'B-JS\n':
                words.append(word)
                i += 1
                if i >= len(t): break
                word = t[i].split('\t')[0]
                tag = t[i].split('\t')[1]
                while tag == 'I-JS\n':
                    words.append(word)
                    i += 1
                    if i >= len(t): break
                    word = t[i].split('\t')[0]
                    tag = t[i].split('\t')[1]
                if ''.join(map(str, words)).lower() not in JS:
                    JS.append(''.join(map(str, words)).lower())
                    words = []
            elif tag == 'B-JR\n':
                words.append(word)
                i += 1
                if i >= len(t): break
                word = t[i].split('\t')[0]
                tag = t[i].split('\t')[1]
                while tag == 'I-JR\n':
                    words.append(word)
                    i += 1
                    if i >= len(t): break
                    word = t[i].split('\t')[0]
                    tag = t[i].split('\t')[1]
                if ''.join(map(str, words)).lower() not in JR:
                    JR.append(''.join(map(str, words)).lower())
                    words = []
            elif tag == 'B-PS\n':
                words.append(word)
                i += 1
                if i >= len(t): break
                word = t[i].split('\t')[0]
                tag = t[i].split('\t')[1]
                while tag == 'I-PS\n':
                    words.append(word)
                    i += 1
                    if i >= len(t): break
                    word = t[i].split('\t')[0]
                    tag = t[i].split('\t')[1]
                if ''.join(map(str, words)).lower() not in PS:
                    PS.append(''.join(map(str, words)).lower())
                    words = []
            else:
                i += 1


        with open(person_predict_bioprocess_path, 'w', encoding='utf-8') as f1:
            f1.write('\t'.join(map(str, JN)))
            f1.write('\n')
            f1.write('\t'.join(map(str, JR)))
            f1.write('\n')
            f1.write('\t'.join(map(str, JS)))
            f1.write('\n')
            f1.write('\t'.join(map(str, PS)))
            f1.write('\n')
            f1.write(education)





    company={}
    #print(company)
    with open(person_information_path,'w',encoding='utf-8') as f2:
        f2.write(repr(company))
        f2.write('\n')
        f2.write(repr(JN))
        f2.write('\n')
        f2.write(repr(JR))
        f2.write('\n')
        f2.write(repr(JS))
        f2.write('\n')
        f2.write(repr(PS))
        f2.write('\n')
        f2.write(text)


    #print(f"cost {time.time()-t_1} total")
    return matchphone,matchemail,matches,matcheslocation,shixi_job,matchesproject,matchesname,matchesgender,matchesbithday,JN,JR,PS,education

if __name__ == '__main__':
    model,device=init_model()
    begin=time.time()
    print(personpredict('./person1/1.pdf',model,device))
    end=time.time()
    print(end-begin)